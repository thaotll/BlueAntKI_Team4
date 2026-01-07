"""
Renders DocxDocumentModel to actual DOCX bytes using python-docx.
Pure rendering logic - no business/content decisions.
"""

import logging
from io import BytesIO
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from app.models.pptx import RgbColor
from app.models.docx import (
    DocxDocumentModel,
    DocxSection,
    DocxParagraph,
    DocxTextRun,
    DocxTextStyle,
    DocxTable,
    DocxTableRow,
    DocxTableCell,
    DocxList,
    DocxListItem,
    DocxHeading,
    HeadingLevel,
    ListStyle,
    TableCellAlignment,
)

logger = logging.getLogger(__name__)


class DocxRenderer:
    """
    Renders document models to DOCX files.

    Usage:
        renderer = DocxRenderer()
        docx_bytes = renderer.render(model)
    """

    # Alignment mapping
    ALIGNMENT_MAP = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    # Heading style mapping
    HEADING_STYLE_MAP = {
        HeadingLevel.H1: "Heading 1",
        HeadingLevel.H2: "Heading 2",
        HeadingLevel.H3: "Heading 3",
        HeadingLevel.H4: "Heading 4",
    }

    def render(self, model: DocxDocumentModel) -> bytes:
        """
        Render document model to DOCX bytes.

        Args:
            model: The document model to render

        Returns:
            DOCX file as bytes
        """
        logger.info(f"Rendering Word document: {model.title}")

        # Create document
        doc = Document()

        # Set document properties
        core_properties = doc.core_properties
        core_properties.title = model.title
        core_properties.author = model.author

        # Render each section
        for i, section_model in enumerate(model.sections):
            logger.debug(f"Rendering section {i + 1}/{len(model.sections)}")
            self._render_section(doc, section_model, model)

        # Save to bytes
        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)

        docx_bytes = stream.getvalue()
        logger.info(f"Rendered {len(model.sections)} sections ({len(docx_bytes)} bytes)")

        return docx_bytes

    def _render_section(
        self,
        doc: Document,
        section_model: DocxSection,
        document_model: DocxDocumentModel,
    ) -> None:
        """Render a single section."""
        # Render heading if present
        if section_model.heading:
            self._render_heading(doc, section_model.heading, document_model)

        # Render content in order if specified
        if section_model.content_order:
            for content_type, index in section_model.content_order:
                if content_type == "paragraph" and index < len(section_model.paragraphs):
                    self._render_paragraph(doc, section_model.paragraphs[index])
                elif content_type == "table" and index < len(section_model.tables):
                    self._render_table(doc, section_model.tables[index], document_model)
                elif content_type == "list" and index < len(section_model.lists):
                    self._render_list(doc, section_model.lists[index])
        else:
            # Default: render paragraphs, then tables, then lists
            for para in section_model.paragraphs:
                self._render_paragraph(doc, para)
            for table in section_model.tables:
                self._render_table(doc, table, document_model)
            for list_model in section_model.lists:
                self._render_list(doc, list_model)

    def _render_heading(
        self,
        doc: Document,
        heading: DocxHeading,
        document_model: DocxDocumentModel,
    ) -> None:
        """Render a heading."""
        style_name = self.HEADING_STYLE_MAP.get(heading.level, "Heading 1")
        para = doc.add_heading(heading.text, level=heading.level.value)

        # Apply custom color if specified
        if heading.color:
            for run in para.runs:
                run.font.color.rgb = self._to_rgb_color(heading.color)

    def _render_paragraph(self, doc: Document, para_model: DocxParagraph) -> None:
        """Render a paragraph with its text runs."""
        para = doc.add_paragraph()
        para.alignment = self.ALIGNMENT_MAP.get(para_model.alignment, WD_ALIGN_PARAGRAPH.LEFT)

        # Set spacing
        para_format = para.paragraph_format
        para_format.space_before = Pt(para_model.space_before_pt)
        para_format.space_after = Pt(para_model.space_after_pt)

        if para_model.first_line_indent_cm > 0:
            para_format.first_line_indent = Cm(para_model.first_line_indent_cm)

        # Render text runs
        for run_model in para_model.runs:
            run = para.add_run(run_model.text)
            if run_model.style:
                self._apply_text_style(run, run_model.style)

    def _render_table(
        self,
        doc: Document,
        table_model: DocxTable,
        document_model: DocxDocumentModel,
    ) -> None:
        """Render a table."""
        if not table_model.rows:
            return

        # Determine number of columns from first row
        num_cols = max(len(row.cells) for row in table_model.rows) if table_model.rows else 0
        if num_cols == 0:
            return

        # Create table
        table = doc.add_table(rows=len(table_model.rows), cols=num_cols)
        table.style = table_model.style_name
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths if specified
        if table_model.col_widths_cm:
            for i, width in enumerate(table_model.col_widths_cm):
                if i < num_cols:
                    for row in table.rows:
                        row.cells[i].width = Cm(width)

        # Populate cells
        for row_idx, row_model in enumerate(table_model.rows):
            row = table.rows[row_idx]
            for col_idx, cell_model in enumerate(row_model.cells):
                if col_idx < num_cols:
                    cell = row.cells[col_idx]
                    self._render_table_cell(cell, cell_model)

    def _render_table_cell(self, cell, cell_model: DocxTableCell) -> None:
        """Render a table cell."""
        # Clear existing paragraphs
        cell.text = ""

        # Add content
        para = cell.paragraphs[0]
        run = para.add_run(cell_model.content)

        # Apply formatting
        if cell_model.bold:
            run.bold = True

        if cell_model.text_color:
            run.font.color.rgb = self._to_rgb_color(cell_model.text_color)

        # Apply alignment
        alignment_map = {
            TableCellAlignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
            TableCellAlignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
            TableCellAlignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
        }
        para.alignment = alignment_map.get(cell_model.alignment, WD_ALIGN_PARAGRAPH.LEFT)

        # Apply background color
        if cell_model.background_color:
            self._set_cell_background(cell, cell_model.background_color)

    def _set_cell_background(self, cell, color: RgbColor) -> None:
        """Set cell background color using XML."""
        hex_color = f"{color.r:02X}{color.g:02X}{color.b:02X}"
        shading_elm = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _render_list(self, doc: Document, list_model: DocxList) -> None:
        """Render a list."""
        for i, item in enumerate(list_model.items):
            if list_model.style == ListStyle.BULLET:
                para = doc.add_paragraph(style="List Bullet")
            else:
                para = doc.add_paragraph(style="List Number")

            run = para.add_run(item.text)
            if item.bold:
                run.bold = True

            # Handle nesting level (simplified - just indent)
            if item.level > 0:
                para.paragraph_format.left_indent = Cm(item.level * 0.5)

    def _apply_text_style(self, run, style: DocxTextStyle) -> None:
        """Apply text styling to a run."""
        run.font.name = style.font_name
        run.font.size = Pt(style.font_size_pt)
        run.bold = style.bold
        run.italic = style.italic
        run.underline = style.underline

        if style.color:
            run.font.color.rgb = self._to_rgb_color(style.color)

    @staticmethod
    def _to_rgb_color(color: RgbColor) -> RGBColor:
        """Convert model RgbColor to python-docx RGBColor."""
        return RGBColor(color.r, color.g, color.b)

